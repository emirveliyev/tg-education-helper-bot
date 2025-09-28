# api/document_generator.py
from typing import Dict, List, Optional
from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from io import BytesIO
from datetime import datetime
import logging

logger = logging.getLogger("tg-edu-bot")

class DocumentGenerator:
    @staticmethod
    def set_doc_defaults(doc: Document) -> None:
        try:
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            style.paragraph_format.space_after = Pt(6)
            style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Arial')
        except Exception as e:
            logger.warning(f"Не удалось установить настройки документа: {e}")

    @staticmethod
    def create_docx_file(meta: Dict, tests: List, header_buf: Optional[BytesIO] = None,
                         out_path: Optional[str] = None, include_answers: bool = False,
                         qtype: str = "closed", wiki_images: Optional[List] = None,
                         wiki_extract: Optional[str] = None) -> Optional[str]:
        try:
            doc = Document()
            DocumentGenerator.set_doc_defaults(doc)

            section = doc.sections[0]
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)

            if header_buf:
                try:
                    header_buf.seek(0)
                    doc.add_picture(header_buf, width=Inches(6.5))
                    doc.add_paragraph()
                except Exception as e:
                    logger.error(f"Ошибка добавления заголовка: {e}")

            title_text = f"{meta.get('subject', 'Предмет')} — {meta.get('topic', 'Тема')}".strip()
            title_par = doc.add_paragraph()
            title_run = title_par.add_run(title_text)
            title_run.bold = True
            title_run.font.size = Pt(20)
            title_par.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            meta_text = (f"Класс: {meta.get('grade', 'Не указан')} | "
                         f"Язык: {meta.get('language', 'Русский')} | "
                         f"Дата: {datetime.now().strftime('%Y-%m-%d')}")
            doc.add_paragraph(meta_text)
            doc.add_paragraph()

            if wiki_extract:
                doc.add_heading("Информация из Википедии", level=2)
                paragraphs = wiki_extract.split("\n\n")
                for para in paragraphs:
                    if para.strip():
                        ex_par = doc.add_paragraph(para.strip())
                        ex_par.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                doc.add_paragraph()

            if wiki_images:
                doc.add_heading("Изображения", level=3)
                for i, imgb in enumerate(wiki_images[:3]):
                    try:
                        imgb.seek(0)
                        doc.add_picture(imgb, width=Inches(4.0))
                        if i < len(wiki_images) - 1:
                            doc.add_paragraph()
                    except Exception as e:
                        logger.error(f"Ошибка добавления изображения {i+1}: {e}")
                doc.add_paragraph()

            if tests:
                doc.add_heading("Вопросы", level=1)

                if qtype == "closed":
                    table = doc.add_table(rows=0, cols=2)
                    table.alignment = WD_TABLE_ALIGNMENT.CENTER
                    table.autofit = True
                    table.allow_autofit = True

                    for idx, test in enumerate(tests, start=1):
                        row = table.add_row()
                        cell_left, cell_right = row.cells

                        target_cell = cell_left if idx % 2 == 1 else cell_right

                        p_question = target_cell.add_paragraph()
                        run_question = p_question.add_run(f"{idx}. {test['question']}")
                        run_question.bold = True

                        options = test.get('options', [])
                        if len(options) == 4:
                            p_options = target_cell.add_paragraph()
                            options_text = f"a) {options[0]}\nb) {options[1]}\nc) {options[2]}\nd) {options[3]}"
                            p_options.add_run(options_text)

                else:
                    for idx, test in enumerate(tests, start=1):
                        p_question = doc.add_paragraph()
                        run_question = p_question.add_run(f"{idx}. {test['question']}")
                        run_question.bold = True

                        doc.add_paragraph("Ответ: " + "_" * 80)
                        doc.add_paragraph()

                doc.add_paragraph()

                if include_answers:
                    doc.add_page_break()
                    doc.add_heading("Ответы (для учителя)", level=1)

                    if qtype == "open":
                        for idx, test in enumerate(tests, start=1):
                            answer = test.get('answer_text', 'Ответ не указан')
                            doc.add_paragraph(f"{idx}. {answer}")
                    else:
                        letters = ['a', 'b', 'c', 'd']
                        for idx, test in enumerate(tests, start=1):
                            answer_idx = test.get('answer', 1)
                            letter = letters[answer_idx - 1] if 1 <= answer_idx <= 4 else "Не указан"
                            doc.add_paragraph(f"{idx}. {letter}")

            if out_path is None:
                ts = datetime.utcnow().strftime("%Y%m%d_%H%M%S")
                out_path = os.path.join(meta.get('config.DATA_DIR', './data'), f"document_{meta.get('user_id', 'unknown')}_{ts}.docx")

            doc.save(out_path)
            logger.info(f"Документ успешно создан: {out_path}")
            return out_path

        except Exception as e:
            logger.error(f"Ошибка создания DOCX файла: {e}")
            return None